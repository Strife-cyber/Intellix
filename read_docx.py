import docx
import os
import sys

files = [
    r'c:\Users\dunam\Desktop\Projects\Personal\Laravel\intellix\storage\app\public\PROSIT ALLER N° 2 Exigence.docx',
    r'c:\Users\dunam\Desktop\Projects\Personal\Laravel\intellix\storage\app\public\Prosit_Aller_1.docx',
    r'c:\Users\dunam\Desktop\Projects\Personal\Laravel\intellix\storage\app\public\Prosit_Aller_4.docx'
]

def extract_text(file_path):
    try:
        doc = docx.Document(file_path)
        fullText = []
        for p in doc.paragraphs:
            if p.text.strip():
                fullText.append(p.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        # replace new lines inside cell with spaces or newlines
                        row_text.append(cell.text.replace("\n", " ").strip())
                if row_text:
                    fullText.append(" | ".join(row_text))
                    
        return '\n'.join(fullText)
    except Exception as e:
        return str(e)

for f in files:
    print(f"=== {os.path.basename(f)} ===")
    print(extract_text(f))
    print("\n" + "="*50 + "\n")
